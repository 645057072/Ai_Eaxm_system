# -*- coding: utf-8 -*-

from datetime import date
from typing import Annotated, List

from fastapi import APIRouter, Depends, File, HTTPException, UploadFile, status
from sqlalchemy import func, select
from sqlalchemy.orm import Session, joinedload

from app.api.deps import get_db, require_permission
from app.core.permissions import is_super_role
from app.core.security import hash_password
from app.models.enterprise import Enterprise
from app.models.student import Student
from app.models.user import Role, User
from app.schemas.common import PageParams, PageResult
from app.schemas.user import UserCreate, UserOut, UserUpdate
from app.services.data_scope import ensure_same_enterprise

router = APIRouter()


def _today() -> date:
    return date.today()


def _parse_date_any(v: str | None) -> date | None:
    if not v:
        return None
    s = v.strip()
    if not s:
        return None
    # 支持 YYYY-MM-DD / YYYY/MM/DD / YYYY.MM.DD
    s = s.replace("/", "-").replace(".", "-")
    try:
        return date.fromisoformat(s)
    except Exception:
        raise HTTPException(status_code=400, detail="日期格式错误，应为 YYYY-MM-DD")


def _ensure_student_linkable(db: Session, current: User, student_id: int, enterprise_id: int | None) -> Student:
    stu = db.get(Student, student_id)
    if stu is None:
        raise HTTPException(status_code=400, detail="关联学员不存在")
    # 学员必须与用户所属企业一致（超管亦如此，避免跨企业错绑）
    if enterprise_id is not None and stu.enterprise_id is not None and stu.enterprise_id != enterprise_id:
        raise HTTPException(status_code=400, detail="学员所属企业与用户所属企业不一致")
    if not is_super_role(current):
        ensure_same_enterprise(current, stu.enterprise_id)
    exists = db.scalar(select(func.count()).select_from(User).where(User.student_id == student_id))
    if exists:
        raise HTTPException(status_code=400, detail="该学员已被其他用户关联")
    return stu


@router.get("", response_model=PageResult[UserOut])
def list_users(
    db: Annotated[Session, Depends(get_db)],
    current: Annotated[User, Depends(require_permission("list.user"))],
    page: Annotated[PageParams, Depends()],
    keyword: str | None = None,
    enterprise_id: int | None = None,
) -> PageResult[UserOut]:
    """用户列表：非超管仅本企业；超管可查全部，可按 enterprise_id 筛选。"""
    stmt = select(func.count()).select_from(User)
    q = select(User).options(joinedload(User.enterprise), joinedload(User.role), joinedload(User.student))
    if is_super_role(current):
        if enterprise_id is not None:
            stmt = stmt.where(User.enterprise_id == enterprise_id)
            q = q.where(User.enterprise_id == enterprise_id)
    else:
        if current.enterprise_id is None:
            return PageResult[UserOut](total=0, items=[])
        stmt = stmt.where(User.enterprise_id == current.enterprise_id)
        q = q.where(User.enterprise_id == current.enterprise_id)
    if keyword:
        stmt = stmt.where(User.username.like(f"%{keyword}%"))
        q = q.where(User.username.like(f"%{keyword}%"))
    total = db.scalar(stmt) or 0
    rows = db.scalars(q.offset(page.skip).limit(page.limit).order_by(User.id.desc())).all()
    return PageResult[UserOut](total=int(total), items=[UserOut.model_validate(r) for r in rows])


@router.post("", response_model=UserOut)
def create_user(
    body: UserCreate,
    db: Annotated[Session, Depends(get_db)],
    current: Annotated[User, Depends(require_permission("action.user.create"))],
) -> UserOut:
    """在本企业创建用户；超管须指定 enterprise_id。"""
    if db.scalar(select(func.count()).select_from(User).where(User.username == body.username)):
        raise HTTPException(status_code=400, detail="用户名已存在")
    role = db.get(Role, body.role_id)
    if role is None:
        raise HTTPException(status_code=400, detail="角色不存在")
    if is_super_role(current):
        if body.enterprise_id is None:
            raise HTTPException(status_code=400, detail="请指定所属企业")
        if db.get(Enterprise, body.enterprise_id) is None:
            raise HTTPException(status_code=400, detail="企业不存在")
        eid = body.enterprise_id
    else:
        if current.enterprise_id is None:
            raise HTTPException(status_code=400, detail="当前账号未关联企业，无法创建用户")
        eid = current.enterprise_id

    if body.student_id is not None:
        _ensure_student_linkable(db, current, body.student_id, eid)

    u = User(
        username=body.username,
        password_hash=hash_password(body.password),
        full_name=body.full_name,
        role_id=body.role_id,
        enterprise_id=eid,
        student_id=body.student_id,
        enable_date=body.enable_date or _today(),
        expire_date=body.expire_date,
    )
    db.add(u)
    db.commit()
    db.refresh(u)
    u = db.scalars(
        select(User)
        .options(joinedload(User.enterprise), joinedload(User.role), joinedload(User.student))
        .where(User.id == u.id)
    ).first()
    return UserOut.model_validate(u)


@router.get("/{user_id}", response_model=UserOut)
def get_user(
    user_id: int,
    db: Annotated[Session, Depends(get_db)],
    current: Annotated[User, Depends(require_permission("list.user"))],
) -> UserOut:
    u = db.scalars(
        select(User).options(joinedload(User.enterprise), joinedload(User.role), joinedload(User.student)).where(User.id == user_id)
    ).first()
    if u is None:
        raise HTTPException(status_code=404, detail="用户不存在")
    if not is_super_role(current):
        ensure_same_enterprise(current, u.enterprise_id)
    return UserOut.model_validate(u)


@router.patch("/{user_id}", response_model=UserOut)
def update_user(
    user_id: int,
    body: UserUpdate,
    db: Annotated[Session, Depends(get_db)],
    current: Annotated[User, Depends(require_permission("action.user.update"))],
) -> UserOut:
    u = db.get(User, user_id)
    if u is None:
        raise HTTPException(status_code=404, detail="用户不存在")
    if not is_super_role(current):
        ensure_same_enterprise(current, u.enterprise_id)
    if body.full_name is not None:
        u.full_name = body.full_name
    if body.role_id is not None:
        if db.get(Role, body.role_id) is None:
            raise HTTPException(status_code=400, detail="角色不存在")
        u.role_id = body.role_id
    if body.is_active is not None:
        u.is_active = body.is_active
    if body.password:
        u.password_hash = hash_password(body.password)
    if body.enable_date is not None:
        u.enable_date = body.enable_date
    if body.expire_date is not None:
        u.expire_date = body.expire_date
    if body.student_id is not None:
        # 允许传 null 取消关联
        if body.student_id:
            _ensure_student_linkable(db, current, body.student_id, u.enterprise_id)
            u.student_id = body.student_id
        else:
            u.student_id = None
    db.commit()
    db.refresh(u)
    u = db.scalars(
        select(User).options(joinedload(User.enterprise), joinedload(User.role), joinedload(User.student)).where(User.id == user_id)
    ).first()
    return UserOut.model_validate(u)


@router.post("/import", response_model=dict)
def import_users(
    file: Annotated[UploadFile, File(...)],
    db: Annotated[Session, Depends(get_db)],
    current: Annotated[User, Depends(require_permission("action.user.import"))],
) -> dict:
    """导入用户：支持 Excel/CSV/TXT/Word(docx)。

    内容字段：用户名、姓名、所属企业、角色、失效日期；启用日期默认导入当天（也可提供）。
    """
    name = (file.filename or "").lower()
    if "." not in name:
        raise HTTPException(status_code=400, detail="文件名缺少扩展名")
    ext = name.rsplit(".", 1)[-1]
    if ext not in ("xls", "xlsx", "csv", "txt", "docx"):
        raise HTTPException(status_code=400, detail="不支持的文件类型")

    raw = file.file.read()

    # 读取为二维表（含表头）
    rows: List[List[str]] = []
    if ext in ("xls", "xlsx"):
        import io

        from openpyxl import load_workbook

        wb = load_workbook(io.BytesIO(raw), data_only=True)
        ws = wb.active
        for r in ws.iter_rows(values_only=True):
            rows.append(["" if v is None else str(v).strip() for v in r])
    elif ext == "docx":
        import io

        from docx import Document

        doc = Document(io.BytesIO(raw))
        if not doc.tables:
            raise HTTPException(status_code=400, detail="Word 文件未找到表格")
        tb = doc.tables[0]
        for tr in tb.rows:
            rows.append([tc.text.strip() for tc in tr.cells])
    else:
        import csv
        import io

        text = raw.decode("utf-8-sig", errors="ignore")
        buf = io.StringIO(text)
        # 尝试自动分隔符（逗号/制表）
        sample = text[:2048]
        dialect = csv.Sniffer().sniff(sample, delimiters=",\t;")
        reader = csv.reader(buf, dialect)
        for r in reader:
            rows.append([c.strip() for c in r])

    if not rows or len(rows) < 2:
        raise HTTPException(status_code=400, detail="导入内容为空")

    header = [h.strip() for h in rows[0]]
    col_map: dict[str, int] = {}
    for i, h in enumerate(header):
        if h in ("用户名", "username"):
            col_map["username"] = i
        elif h in ("姓名", "full_name", "name"):
            col_map["full_name"] = i
        elif h in ("所属企业", "enterprise", "enterprise_name"):
            col_map["enterprise"] = i
        elif h in ("角色", "role", "role_name"):
            col_map["role"] = i
        elif h in ("失效日期", "expire_date"):
            col_map["expire_date"] = i
        elif h in ("启用日期", "enable_date"):
            col_map["enable_date"] = i

    for k in ("username", "full_name", "enterprise", "role"):
        if k not in col_map:
            raise HTTPException(status_code=400, detail="导入表头缺少必填列：用户名/姓名/所属企业/角色")

    created = 0
    skipped = 0
    for r in rows[1:]:
        if not any((c or "").strip() for c in r):
            continue
        username = (r[col_map["username"]] if col_map["username"] < len(r) else "").strip()
        full_name = (r[col_map["full_name"]] if col_map["full_name"] < len(r) else "").strip()
        ent_name = (r[col_map["enterprise"]] if col_map["enterprise"] < len(r) else "").strip()
        role_name = (r[col_map["role"]] if col_map["role"] < len(r) else "").strip()
        expire_s = (r[col_map["expire_date"]] if col_map.get("expire_date", -1) < len(r) else "").strip() if "expire_date" in col_map else ""
        enable_s = (r[col_map["enable_date"]] if col_map.get("enable_date", -1) < len(r) else "").strip() if "enable_date" in col_map else ""

        if not username or not full_name or not ent_name or not role_name:
            skipped += 1
            continue

        if db.scalar(select(func.count()).select_from(User).where(User.username == username)):
            skipped += 1
            continue

        ent = db.scalars(select(Enterprise).where(Enterprise.name == ent_name)).first()
        if ent is None:
            skipped += 1
            continue

        # 普通用户导入仅能导入本企业
        if not is_super_role(current):
            ensure_same_enterprise(current, ent.id)

        role = db.scalars(select(Role).where(Role.name == role_name)).first()
        if role is None:
            skipped += 1
            continue

        enable_d = _parse_date_any(enable_s) or _today()
        expire_d = _parse_date_any(expire_s)

        # 密码策略：导入用户默认密码=用户名（避免空密码）
        u = User(
            username=username,
            password_hash=hash_password(username),
            full_name=full_name,
            role_id=role.id,
            enterprise_id=ent.id,
            enable_date=enable_d,
            expire_date=expire_d,
        )
        db.add(u)
        created += 1

    db.commit()
    return {"created": created, "skipped": skipped}


@router.delete("/{user_id}", status_code=status.HTTP_204_NO_CONTENT)
def delete_user(
    user_id: int,
    db: Annotated[Session, Depends(get_db)],
    current: Annotated[User, Depends(require_permission("action.user.delete"))],
) -> None:
    if user_id == current.id:
        raise HTTPException(status_code=400, detail="不能删除当前登录用户")
    u = db.get(User, user_id)
    if u is None:
        raise HTTPException(status_code=404, detail="用户不存在")
    if not is_super_role(current):
        ensure_same_enterprise(current, u.enterprise_id)
    db.delete(u)
    db.commit()
